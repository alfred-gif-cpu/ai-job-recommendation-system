"""Shared pytest fixtures.

Resume fixtures are generated in-memory (not committed as binary files) so
they always match what the tests expect and never go stale.
"""

import io
import os
import sys

import pytest

sys.path.insert(0, os.path.dirname(os.path.dirname(__file__)))

# Match the app's runtime: classic matcher, real Adzuna keys available so the
# live-search tests can exercise real request-building code paths (network
# calls themselves are mocked, never hit in tests).
os.environ.setdefault("USE_SEMANTIC", "0")
os.environ.setdefault("ADZUNA_APP_ID", "test_app_id")
os.environ.setdefault("ADZUNA_APP_KEY", "test_app_key")


@pytest.fixture(scope="session")
def app_module():
    import app as app_module
    return app_module


@pytest.fixture
def client(app_module):
    app_module.app.config["TESTING"] = True
    return app_module.app.test_client()


@pytest.fixture
def pdf_resume_bytes():
    """A simple one-page PDF resume as an in-memory PDF file."""
    from fpdf import FPDF

    text = (
        "RIYA SHARMA - Data Scientist\n"
        "Skills: Python, Machine Learning, Deep Learning, TensorFlow, PyTorch, "
        "Pandas, NumPy, SQL, NLP.\n"
        "Built recommendation systems and deployed models with Flask on AWS."
    )
    pdf = FPDF()
    pdf.set_margins(15, 15, 15)
    pdf.add_page()
    pdf.set_font("Helvetica", size=12)
    pdf.multi_cell(w=pdf.epw, h=8, text=text)
    return bytes(pdf.output())


@pytest.fixture
def docx_resume_bytes():
    """A simple .docx resume (with a table) as in-memory bytes."""
    import docx

    d = docx.Document()
    d.add_heading("Riya Sharma - Data Scientist", level=1)
    d.add_paragraph(
        "Skills: Python, Machine Learning, Deep Learning, TensorFlow, "
        "PyTorch, Pandas, NumPy, SQL, NLP."
    )
    d.add_paragraph("Built recommendation systems and deployed models with Flask on AWS.")
    table = d.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Kubernetes"
    table.rows[0].cells[1].text = "Docker"

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


@pytest.fixture
def sample_adzuna_response():
    """A realistic Adzuna API JSON payload for mocking urlopen."""
    return {
        "results": [
            {
                "title": "Python Machine Learning Developer",
                "company": {"display_name": "Acme India"},
                "location": {"display_name": "Bangalore, Karnataka"},
                "description": "We need <strong>python</strong>, sql, aws and machine learning experience.",
                "redirect_url": "https://www.adzuna.in/details/1",
                "salary_min": 800000,
                "salary_max": 1800000,
            },
            {
                "title": "Data Analyst",
                "company": {"display_name": "DataCorp"},
                "location": {"display_name": "Pune"},
                "description": "excel &amp; tableau reporting role, work from home available",
                "redirect_url": "https://www.adzuna.in/details/2",
                "salary_min": 500000,
                "salary_max": 500000,
                "salary_is_predicted": "1",
            },
            {
                "title": "Remote Sensing Engineer",
                "company": {"display_name": "GeoSpatial Labs"},
                "location": {"display_name": "Hyderabad"},
                "description": "satellite remote sensing and GIS analysis, in-office role",
                "redirect_url": "https://www.adzuna.in/details/3",
            },
        ]
    }
