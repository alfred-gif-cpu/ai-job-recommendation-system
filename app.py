import os


def _load_env_file(path=".env"):
    """Load KEY=VALUE lines from a .env file into the environment.

    Real environment variables always win (we only set defaults), and there is
    no external dependency. Called before anything reads the Adzuna keys.
    """
    if not os.path.exists(path):
        return
    with open(path, encoding="utf-8") as f:
        for line in f:
            line = line.strip()
            if not line or line.startswith("#") or "=" not in line:
                continue
            key, _, val = line.partition("=")
            os.environ.setdefault(key.strip(), val.strip().strip('"').strip("'"))


_load_env_file()

from flask import Flask, jsonify, render_template, request
from smart_india_recommender import recommend_jobs
from job_api import fetch_live_jobs, INDIA_CITIES
from skills_config import extract_skills, learn_link
import PyPDF2
import docx

app = Flask(__name__)

# Reject oversized uploads (16 MB) instead of buffering huge files.
app.config["MAX_CONTENT_LENGTH"] = 16 * 1024 * 1024
app.jinja_env.globals["learn_link"] = learn_link

RESUME_EXTENSIONS = (".pdf", ".docx")


# -----------------------------
# RESUME TEXT EXTRACTION
# -----------------------------
def extract_text_from_pdf(pdf_file):
    text = ""
    reader = PyPDF2.PdfReader(pdf_file)

    for page in reader.pages:
        # extract_text() can return None for image-only / scanned pages
        text += (page.extract_text() or "") + " "

    return text.lower()


def extract_text_from_docx(docx_file):
    document = docx.Document(docx_file)

    parts = [p.text for p in document.paragraphs if p.text]
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)

    return " ".join(parts).lower()


def extract_text_from_resume(resume_file, filename):
    """Dispatch to the right parser based on file extension."""
    if filename.lower().endswith(".docx"):
        return extract_text_from_docx(resume_file)
    return extract_text_from_pdf(resume_file)


def dedupe_skills(raw):
    """Split a comma string into unique skills, preserving order."""
    seen = set()
    out = []
    for s in raw.split(","):
        s = s.strip()
        if s and s.lower() not in seen:
            seen.add(s.lower())
            out.append(s)
    return out


def find_jobs(skills, live, where=""):
    """Return (results, message, live_ok). Uses Adzuna in live mode, else CSV.

    Falls back to the local CSV recommender when live search is unavailable so
    the app always returns something useful. ``live_ok`` is True only when the
    results genuinely came from the live API.
    """
    if live:
        results, err = fetch_live_jobs(skills, where=where)
        if err:
            return recommend_jobs(skills), err + " Showing sample matches instead.", False
        return results, "", True
    return recommend_jobs(skills), "", False


@app.route("/", methods=["GET", "POST"])
def home():

    results = []
    detected_skills = []
    skills_input = ""
    message = ""
    resume_handled = False
    live_ok = False

    if request.method == "POST":
        skills_input = (request.form.get("skills") or "").strip()
        live = request.form.get("live") in ("on", "1")
        where = (request.form.get("where") or "").strip()
        resume = request.files.get("resume")

        # Resume upload flow
        if resume and resume.filename:
            resume_handled = True
            if not resume.filename.lower().endswith(RESUME_EXTENSIONS):
                message = "Please upload a PDF or Word (.docx) file."
            else:
                try:
                    text = extract_text_from_resume(resume, resume.filename)
                except Exception:
                    text = ""
                    message = "Sorry, we couldn't read that file. Try another one."

                detected_skills = extract_skills(text)

                if not detected_skills and not message:
                    message = (
                        "No known skills were detected in that resume. "
                        "Try typing your skills instead."
                    )
                elif detected_skills:
                    # Carry the detected skills into the text box so the user can
                    # change the city / re-search without re-uploading the resume.
                    skills_input = ", ".join(detected_skills)
                    results, message, live_ok = find_jobs(skills_input, live, where)
    else:
        # GET: support shareable links like /?skills=python,sql&live=1&where=Pune
        skills_input = (request.args.get("skills") or "").strip()
        live = request.args.get("live") in ("on", "1")
        where = (request.args.get("where") or "").strip()

    # Manual / shared-link skills flow
    if not resume_handled and skills_input:
        detected_skills = dedupe_skills(skills_input)
        results, message, live_ok = find_jobs(skills_input, live, where)
    elif not resume_handled and request.method == "POST":
        message = "Please enter your skills or upload a resume."

    # Ran a search but nothing matched
    if not results and not message and (detected_skills or skills_input):
        message = "No matching jobs found. Try different or more common skills."

    return render_template(
        "index.html",
        results=results,
        detected_skills=detected_skills,
        skills_input=skills_input,
        message=message,
        live_requested=live,
        live_on=live_ok,
        where=where,
        cities=INDIA_CITIES,
    )


@app.route("/live-more")
def live_more():
    """Fetch one more page of live Adzuna jobs (on-demand pagination).

    Returns JSON with rendered card HTML and whether more pages likely exist.
    """
    skills = (request.args.get("skills") or "").strip()
    where = (request.args.get("where") or "").strip()
    try:
        page = int(request.args.get("page", 2))
    except ValueError:
        page = 2

    if not skills or page < 2:
        return jsonify(html="", count=0, hasMore=False)

    jobs, err = fetch_live_jobs(skills, where=where, page=page)
    if err:
        return jsonify(html="", count=0, hasMore=False, error=err)

    html = render_template("_cards.html", results=jobs)
    # A full page suggests there may be another; a short page means we're done.
    return jsonify(html=html, count=len(jobs), hasMore=len(jobs) >= 50)


if __name__ == "__main__":
    app.run(host="0.0.0.0", port=10000)
